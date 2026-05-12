"""Generate rewritten Curriculum Fundamentals docx organized by defense layers and 4-phase structure."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def build_doc():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ══════════════════════════════════════════════════════════════
    # TITLE
    # ══════════════════════════════════════════════════════════════
    title = doc.add_heading('Fundamentals Curriculum', level=0)
    add_para(doc, 'A positional framework for teaching Brazilian Jiu-Jitsu fundamentals.', italic=True)

    # ══════════════════════════════════════════════════════════════
    # ABSTRACT
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'Abstract', 1)

    doc.add_paragraph(
        'This document reorganises the fundamentals curriculum around two structural axes: '
        'the defense-layer hierarchy and the four-phase positional vector. '
        'Together they provide a coordinate system that locates every technique within a single, '
        'coherent map of the game.'
    )

    add_heading(doc, 'Defense Layers', 2)
    doc.add_paragraph(
        'Guard players defend through successive barriers. Each layer cleared by the passer '
        'brings them closer to a dominant position. The layers, from outermost to innermost:'
    )
    add_bullets(doc, [
        'Level 3 — Feet (Open Guard): First line of defense. Feet and legs create distance, hooks, and entanglements. '
        'Includes collar sleeve guard, de la Riva, and other open guard variations.',
        'Level 2 — Knees (Half Guard): Feet cleared. One knee entangles the opponent\'s leg.',
        'Level 1 — Hips (Closed Guard): Last barrier. Both legs wrap the torso, feet locked.',
        'Level 0 — Passed: Guard cleared. Dominant positions — side control, mount, back control, turtle.',
    ])
    doc.add_paragraph(
        'A pass advances through the layers (3→2→1→0). An escape recovers them in reverse (0→1→2→3). '
        'Every technique in the curriculum lives at one of these levels.'
    )

    add_heading(doc, 'The Four-Phase Vector', 2)
    doc.add_paragraph(
        'For every position at every defense layer, the curriculum addresses four phases:'
    )
    add_bullets(doc, [
        'Define — What structurally identifies this position. The constraints that make it what it is.',
        'Control — How to establish, maintain, and stabilise the position. Structure, grips, pressure, balance.',
        'Attack / Defend — What can be done from the position. '
        'Attack (Me): submissions, sweeps, advances. Defend (Op): survival, frames, escaping submissions.',
        'Transit / React — How to move to the next position. '
        'Transit (Me): sweeps, guard recovery, positional advance. React (Op): responding to the opponent\'s movement.',
    ])
    doc.add_paragraph(
        'The "Me" and "Op" perspectives correspond to playing the position versus opposing it. '
        'For example, closed guard Me = the guard player; closed guard Op = the passer. '
        'This dual perspective is central: students learn both sides of every position, '
        'and understanding one side deepens understanding of the other.'
    )

    # ══════════════════════════════════════════════════════════════
    # PHILOSOPHY
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'Curriculum Philosophy', 1)

    doc.add_paragraph(
        'The idea of the curriculum is not only to teach techniques, but to help students understand:'
    )
    add_bullets(doc, [
        'What defines the position',
        'How to perform and maintain the position',
        'What can be done from the position',
    ])

    doc.add_paragraph('For every position a student should understand:')

    add_heading(doc, 'Identity of the Position', 3)
    doc.add_paragraph('What defines the position and how it is called.')
    doc.add_paragraph('Examples:')
    add_bullets(doc, [
        'Closed guard is when the person on bottom holds the person on top by crossing the feet around the waist.',
        'De la Riva is one outside hook on the leg of the person on top.',
        'Mount is when someone is sitting on the torso of someone with both knees on the floor.',
    ])

    add_heading(doc, 'Structure', 3)
    doc.add_paragraph('How to perform and stabilise the position, how to make it efficient.')
    doc.add_paragraph('Examples:')
    add_bullets(doc, [
        'Closed guard: hold one collar to control posture.',
        'De la Riva: use one grip on the pants, or on the sleeve, to make the position more efficient.',
    ])

    add_heading(doc, 'Actions From the Position', 3)
    doc.add_paragraph('What the student can do from there:')
    add_bullets(doc, [
        'Attacks and submissions',
        'Sweeps',
        'Transitions',
        'Escapes and reactions',
    ])

    doc.add_paragraph(
        'The goal is for students to understand positions as systems rather than isolated techniques.'
    )

    add_heading(doc, 'Guiding Principles', 2)
    doc.add_paragraph('The curriculum follows the sequence of a match and prioritises:')
    add_bullets(doc, [
        'Control before attack',
        'Survival before escape',
        'Posture before movement',
        'Understanding before memorisation',
    ])

    doc.add_paragraph('Each class is designed around:')
    add_bullets(doc, [
        'One main positional theme',
        'Specific goals',
        'Progressive resistance',
        'Positional sparring',
        'Live application',
    ])

    doc.add_paragraph('Students are also informed about:')
    add_bullets(doc, [
        'What they are learning during the week',
        'What is coming next',
        'The main goal of the class',
        'The key concept they should understand',
    ])

    add_heading(doc, 'Attack and Defense Cycles', 2)
    doc.add_paragraph(
        'Another important principle of the curriculum structure is that each main theme or position '
        'is generally taught in attacking and defensive cycles, usually over 2-week blocks each.'
    )
    doc.add_paragraph('The progression normally follows:')
    add_bullets(doc, [
        'Understanding and controlling the position',
        'Attacking and advancing from the position',
        'Defending and surviving the same position',
        'Escaping and recovering from the same position',
    ])
    doc.add_paragraph('Example:')
    add_bullets(doc, [
        'Side control top → control, pressure, transitions and attacks',
        'Side control bottom → frames, survival, escapes and recovery',
    ])
    doc.add_paragraph(
        'This structure allows students to reinforce what they learned previously from the opposite perspective.'
    )
    doc.add_paragraph(
        'When students already understand what the top player is trying to do, '
        'they begin to better understand how to defend and escape the position.'
    )
    doc.add_paragraph(
        'Likewise, after learning defence and escapes, students return to the attacking side '
        'with a much deeper understanding of timing, reactions, pressure, openings, and defensive structures.'
    )
    doc.add_paragraph('This creates continuous reinforcement between:')
    add_bullets(doc, [
        'Attack and defence',
        'Top and bottom',
        'Control and escape',
    ])
    doc.add_paragraph(
        'The objective is for students to understand the complete relationship of the position '
        'rather than isolated techniques.'
    )
    doc.add_paragraph('This progression also helps:')
    add_bullets(doc, [
        'Reduce confusion',
        'Improve retention of information',
        'Reinforce previous learning',
        'Develop problem-solving skills',
        'Build confidence progressively',
    ])
    doc.add_paragraph(
        'By revisiting the same positions from different perspectives, students develop '
        'a much stronger conceptual understanding of Jiu-Jitsu.'
    )

    # ══════════════════════════════════════════════════════════════
    # LEVEL 3: FEET — OPEN GUARD
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'Level 3 — Feet: Open Guard', 1)

    # --- DEFINE ---
    add_heading(doc, 'Define', 2)
    doc.add_paragraph(
        'Open guard is the first line of defense. The guard player uses feet, legs, and grips '
        'to create distance, control the passer, and prevent advancement to closer positions.'
    )

    add_heading(doc, 'Collar Sleeve Guard — Position Identity', 3)
    doc.add_paragraph('The objective of collar sleeve guard is to:')
    add_bullets(doc, [
        'Control distance',
        'Control posture and balance',
        'Manage space using grips, hips and legs',
        'Create angles and off-balancing opportunities',
    ])
    doc.add_paragraph(
        'Students should understand that collar sleeve is an open guard position used to '
        'prevent pressure, maintain connection, create movement, and transition to stronger positions or attacks.'
    )

    add_heading(doc, 'Guard Retention — Position Identity', 3)
    doc.add_paragraph('Students were introduced to the idea that guard retention is based on:')
    add_bullets(doc, [
        'Frames',
        'Hip mobility',
        'Angle management',
        'Distance management',
        'Reconnecting structure',
    ])
    doc.add_paragraph(
        'The goal was for students to understand that retention is not only using the legs, '
        'but coordinating hips, knees, frames, head position, and grips to recover and maintain the guard.'
    )

    # --- CONTROL ---
    add_heading(doc, 'Control', 2)
    doc.add_paragraph(
        'Guard is not only attacking — it is also maintaining structure, recovering position, '
        'managing distance, preventing passes, and reconnecting hips and legs.'
    )

    add_heading(doc, 'Collar Sleeve — Entries', 3)
    doc.add_paragraph('Students learn how to enter collar sleeve guard:')
    add_bullets(doc, [
        'From standing situations',
        'From seated / sit-up guard situations',
    ])
    doc.add_paragraph('Main concepts:')
    add_bullets(doc, [
        'Establishing grips early',
        'Controlling distance',
        'Connecting hips and legs before opponent closes space',
    ])

    add_heading(doc, 'Collar Sleeve — Retention and Maintenance', 3)
    doc.add_paragraph('Main focus:')
    add_bullets(doc, [
        'Maintaining guard structure',
        'Recovering distance',
        'Reconnecting legs and hips when pressure is applied',
    ])
    doc.add_paragraph('Concepts taught:')
    add_bullets(doc, [
        'Hip follow',
        'Leg pummelling',
        'Shallow lasso',
    ])
    doc.add_paragraph(
        'The goal is for students to understand how to stay connected, '
        'how to prevent the passer from controlling the hips, '
        'and how to recover structure when the guard begins to break.'
    )

    add_heading(doc, 'Guard Retention — Beginner Progression', 3)
    doc.add_paragraph(
        'The progression started very basic and gradually increased in pace and pressure through the week.'
    )
    doc.add_paragraph('For beginners the training started slowly and with minimal pressure.')
    doc.add_paragraph('Main focus:')
    add_bullets(doc, [
        'Recovering guard when opponent passes on the knees',
        'Replacing guard using the knee',
        'Controlling the cross face',
        'Squaring the hips',
        'Moving the head in the opposite direction of the pass',
    ])
    doc.add_paragraph('Once students became more comfortable, progression moved into:')
    add_bullets(doc, [
        'Toreando-style passing situations',
        'Recovering angles and reconnecting the legs',
    ])
    doc.add_paragraph(
        'The emphasis remained on movement quality, positioning, and understanding '
        'the role of the frames and hips rather than speed or athleticism.'
    )

    # --- ATTACK / DEFEND ---
    add_heading(doc, 'Attack / Defend', 2)

    add_heading(doc, 'Attack (Me) — Actions From Collar Sleeve Guard', 3)
    doc.add_paragraph(
        'Once students understand the position and how to maintain it, they begin learning '
        'the main actions available from collar sleeve guard.'
    )
    doc.add_paragraph('Main actions:')
    add_bullets(doc, [
        'Transition to closed guard',
        'Tripod sweep',
    ])
    doc.add_paragraph(
        'The objective is not simply to perform techniques, but to understand '
        'when to off-balance, when to reconnect, when to transition, and when to sweep.'
    )
    doc.add_paragraph(
        'Students should begin understanding collar sleeve guard as a position that creates '
        'control, movement, transitions, and imbalance.'
    )

    add_heading(doc, 'Attack (Me) — Guard Retention to Offense', 3)
    doc.add_paragraph(
        'Once students understood the defensive structure of retention, they began connecting retention to:'
    )
    add_bullets(doc, [
        'Re-establishing guard',
        'Transitioning back to collar sleeve',
        'Recovering closed guard',
        'Creating off-balancing opportunities',
    ])
    doc.add_paragraph('The objective of the block was to build early understanding that:')
    add_bullets(doc, [
        'Guard retention is part of guard offence',
        'Maintaining structure creates attacking opportunities',
        'Recovering position is often more important than forcing attacks',
    ])

    add_heading(doc, 'Defend (Op) — Guard Passing', 3)
    doc.add_paragraph(
        'After working on collar sleeve guard, guard retention, transitions to closed guard '
        'and closed guard attacks, the curriculum progressed into the top game and guard passing.'
    )
    doc.add_paragraph('The objective of this phase was to help students understand:')
    add_bullets(doc, [
        'How to establish safe posture inside the guard',
        'How to open the guard correctly',
        'How to maintain the guard open',
        'How to control before passing',
        'How passing connects to previously learned guards',
    ])
    doc.add_paragraph(
        'The progression was designed so students could now begin understanding '
        'the opposite side of the positions they had already learned from the bottom.'
    )

    add_heading(doc, 'Guard Passing — Position Identity (Passer\'s Perspective)', 4)
    doc.add_paragraph('Students were introduced to the idea that guard passing is based on:')
    add_bullets(doc, [
        'Posture',
        'Balance',
        'Base',
        'Pressure',
        'Controlling distance',
        'Disconnecting the guard before passing',
    ])
    doc.add_paragraph('The main objective of the passer is to:')
    add_bullets(doc, [
        'Stay safe',
        'Avoid being off-balanced',
        'Open the guard',
        'Control the hips and legs',
        'Progress to dominant positions',
    ])
    doc.add_paragraph('Students were encouraged to understand that:')
    add_bullets(doc, [
        'Posture comes before opening',
        'Opening comes before passing',
        'Control comes before movement',
    ])

    add_heading(doc, 'Guard Passing — Structure and Execution', 4)
    doc.add_paragraph('The progression started with:')
    add_bullets(doc, [
        'Base inside closed guard',
        'Grips and posture',
        'Understanding alignment and balance',
    ])
    doc.add_paragraph('Before moving into:')
    add_bullets(doc, [
        'Opening the guard',
        'Maintaining the guard open',
        'Beginning the passing process',
    ])

    add_heading(doc, 'Week 1 — Passing from the Knees', 4)
    doc.add_paragraph('Main focus:')
    add_bullets(doc, [
        'Strong base and posture',
        'Safe guard opening from the knees',
        'Maintaining structure while opening',
        'Controlling the legs after the guard opens',
        'Beginning passing mechanics from kneeling positions',
    ])
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'How to stay balanced while opening the guard',
        'How to avoid posture being broken',
        'How to maintain pressure and connection while passing',
    ])

    add_heading(doc, 'Week 2 — Standing Passing', 4)
    doc.add_paragraph(
        'Once students understood posture and opening mechanics from the knees, '
        'progression moved into standing passing.'
    )
    doc.add_paragraph('Main focus:')
    add_bullets(doc, [
        'Posture and grips while standing',
        'Opening the guard from standing',
        'Maintaining the guard open',
        'Toreando-style movement and passing',
        'Passing collar sleeve situations',
    ])
    doc.add_paragraph(
        'This also reinforced concepts from previous blocks, helping students connect '
        'guard retention, collar sleeve guard, distance management, and passing reactions.'
    )
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'How movement creates passing opportunities',
        'How disconnecting the legs creates space',
        'How controlling angles and distance affects passing success',
    ])

    add_heading(doc, 'Guard Passing — Key Sequence', 4)
    doc.add_paragraph(
        'The objective of the block was to help students understand guard passing as '
        'a sequence, a process of control and progression — not simply forcing movement or rushing to pass.'
    )
    doc.add_paragraph('Key concept: posture → open → control → pass')

    # --- TRANSIT / REACT ---
    add_heading(doc, 'Transit / React', 2)
    doc.add_paragraph(
        'From open guard, the main transitions are closing the guard (OGRD → CGRD) '
        'and recovering guard from inferior positions. '
        'The passer reacts to sweeps and retention by adjusting base, opening the guard, '
        'and progressing through the defense layers.'
    )

    # ══════════════════════════════════════════════════════════════
    # LEVEL 1: HIPS — CLOSED GUARD
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'Level 1 — Hips: Closed Guard', 1)

    # --- DEFINE ---
    add_heading(doc, 'Define', 2)
    doc.add_paragraph(
        'After working on collar sleeve guard, guard retention, tripod sweep and transitions '
        'into closed guard, the curriculum progressed into closed guard attacks.'
    )
    doc.add_paragraph('The objective of this phase was to help students understand:')
    add_bullets(doc, [
        'How to establish closed guard',
        'How to control posture',
        'How to create attacks from a stable position',
        'How to connect sweeps and submissions together',
    ])
    doc.add_paragraph('The transition into closed guard usually started from:')
    add_bullets(doc, [
        'Pulling collar sleeve',
        'Creating connection',
        'Closing the guard before attacking',
    ])

    add_heading(doc, 'Position Identity', 3)
    doc.add_paragraph('Students were introduced to the idea that closed guard is a control position based on:')
    add_bullets(doc, [
        'Posture control',
        'Connection with the legs',
        'Breaking alignment',
        'Controlling distance and balance',
    ])
    doc.add_paragraph('The main objective of the position is to:')
    add_bullets(doc, [
        'Break posture',
        'Control movement',
        'Create attacking opportunities safely',
    ])
    doc.add_paragraph(
        'Students were encouraged to understand that before attacking: '
        'posture must be controlled, angles must be created, and structure must be maintained.'
    )

    # --- CONTROL ---
    add_heading(doc, 'Control', 2)
    doc.add_paragraph('Main focus:')
    add_bullets(doc, [
        'Pulling collar sleeve and establishing closed guard',
        'Breaking posture using grips and legs',
        'Maintaining connection and control',
        'Understanding how posture affects attacks and sweeps',
    ])
    doc.add_paragraph('The progression started with:')
    add_bullets(doc, [
        'Posture breaking',
        'Collar control',
        'Controlling the arms and alignment',
    ])
    doc.add_paragraph('Before moving into attacks.')
    doc.add_paragraph('Students were taught that:')
    add_bullets(doc, [
        'Controlling posture is more important than rushing submissions',
        'Losing posture control usually means losing attacking opportunities',
    ])

    # --- ATTACK / DEFEND ---
    add_heading(doc, 'Attack / Defend', 2)

    add_heading(doc, 'Attack (Me)', 3)
    doc.add_paragraph(
        'Once posture and control were established, students began working on:'
    )
    add_bullets(doc, [
        'Cross collar choke',
        'Scissor sweep',
    ])
    doc.add_paragraph('The objective was for students to begin understanding:')
    add_bullets(doc, [
        'How attacks connect together',
        'How posture reactions create opportunities',
        'How sweeps and submissions work together',
    ])
    doc.add_paragraph(
        'Students were encouraged to see closed guard as a position of control first, '
        'then a platform for attacks and transitions.'
    )

    # --- TRANSIT / REACT ---
    add_heading(doc, 'Transit / React', 2)
    doc.add_paragraph(
        'From closed guard, the main transitions are opening the guard (passer\'s objective) '
        'and the guard player\'s counter-transitions. Guard passing from closed guard connects '
        'to the open guard concepts already learned — once the guard opens, the game returns to Level 3.'
    )

    # ══════════════════════════════════════════════════════════════
    # LEVEL 0: PASSED — SIDE CONTROL
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'Level 0 — Passed: Side Control', 1)

    doc.add_paragraph(
        'After working on collar sleeve guard (entries, retention, attacks and transitions to closed guard), '
        'followed by closed guard (posture breaking, control and attacks), and closed guard opening and passing, '
        'the curriculum progressed into side control.'
    )
    doc.add_paragraph('The objective of this phase was to help students understand:')
    add_bullets(doc, [
        'What defines strong side control',
        'How to establish and maintain the position',
        'How to use body positioning and pressure correctly',
        'How to transition between dominant positions',
        'How control creates attacking opportunities',
    ])
    doc.add_paragraph(
        'The progression was designed so students could now begin understanding '
        'how to stabilise dominant positions after the pass.'
    )

    # --- DEFINE ---
    add_heading(doc, 'Define', 2)
    doc.add_paragraph('Students were introduced to the idea that side control is based on:')
    add_bullets(doc, [
        'Chest-to-chest connection',
        'Head and hip control',
        'Controlling movement and space',
        'Pressure and weight distribution',
        'Balance and base',
    ])
    doc.add_paragraph('The main objective of side control is to:')
    add_bullets(doc, [
        'Limit movement',
        'Control the hips and shoulders',
        'Maintain dominant positioning',
        'Create openings for attacks and transitions',
    ])
    doc.add_paragraph('Students were encouraged to understand that:')
    add_bullets(doc, [
        'Control comes before attack',
        'Maintaining position is more important than rushing submissions',
        'Body positioning and pressure are more important than strength',
    ])

    # --- CONTROL (Me) ---
    add_heading(doc, 'Control (Me) — Side Control Top', 2)

    doc.add_paragraph('The progression started with:')
    add_bullets(doc, [
        'Understanding body positioning',
        'Chest-to-chest control',
        'Head and hip control',
        'Pins and pressure',
        'Base and balance',
    ])
    doc.add_paragraph('Before moving into:')
    add_bullets(doc, [
        'Switching bases',
        'Maintaining control under movement and resistance',
        'Transitioning between control positions',
    ])
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'How to follow movement while maintaining pressure',
        'How base changes affect control',
        'How positioning affects stability and attacking opportunities',
    ])

    add_heading(doc, 'Week 1 — Establishing and Maintaining Side Control', 3)
    doc.add_paragraph('Main focus: body positioning.')
    doc.add_paragraph(
        'Ask the students to do a side control specific without giving much explanation (2 minutes each max). '
        'After, ask who managed to hold the partner without losing the position.'
    )
    doc.add_paragraph(
        'Now explain the use of the hands is not mainly to hold the partner, '
        'but to assist the pressure control along with your body position, '
        'block escapes and set up transitions and submissions.'
    )
    doc.add_paragraph(
        'Now ask them to hold the partner without hands only using the body and explain briefly '
        'how to pin, use the foot and legs and changing bases (2 minutes each) low intensity 30% max.'
    )
    doc.add_paragraph(
        'Next ask to use the hand that controls the hips and explain the use of this hand '
        '(2 minutes each) low intensity.'
    )
    doc.add_paragraph(
        'Now ask to use the hand that controls the head (strongest control) '
        'and same intensity and time.'
    )
    doc.add_paragraph(
        'Finale: do another specific from side control with more focus in controlling the position '
        'with higher intensity. They should improve their side control.'
    )
    doc.add_paragraph(
        'Once students became more comfortable maintaining the position, progression moved into '
        'basic transitions to mount and simple submissions when control was stable.'
    )
    doc.add_paragraph('The emphasis remained on:')
    add_bullets(doc, [
        'Maintaining structure',
        'Balance and positioning',
        'Controlling movement before attacking',
    ])

    add_heading(doc, 'Week 2 — Control to Progression', 3)
    doc.add_paragraph(
        'Once students understood how to stabilise side control, '
        'progression moved into attacks and positional advancement.'
    )
    doc.add_paragraph('Main focus:')
    add_bullets(doc, [
        'Maintaining and adjusting side control',
        'Recognising attacking opportunities',
        'Transitioning between control and attack',
        'Understanding reactions from the bottom player',
    ])
    doc.add_paragraph('Main actions:')
    add_bullets(doc, [
        '1 submission (americana or kimura)',
        '1 transition to mount',
    ])
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'How attacks come from control',
        'How reactions create openings',
        'How transitions and submissions connect together',
    ])

    # --- ATTACK (Me) ---
    add_heading(doc, 'Attack (Me) — Actions From Side Control', 2)
    doc.add_paragraph(
        'Once students understood the structure and control of side control, '
        'they began connecting the position to:'
    )
    add_bullets(doc, [
        'Submissions',
        'Mount transitions',
        'Maintaining pressure under movement',
        'Positional progression',
    ])
    doc.add_paragraph(
        'The objective of the block was to help students understand side control as '
        'a position of control first, then a platform for attacks and transitions.'
    )
    doc.add_paragraph('Key concepts:')
    add_bullets(doc, [
        'Control before attack',
        'Pressure before movement',
        'Maintain position before advancing',
    ])

    # --- DEFEND (Op) — Side Control Defence ---
    add_heading(doc, 'Defend (Op) — Side Control Defence', 2)
    doc.add_paragraph(
        'Now that students had started to understand control and attacks from top side control, '
        'the curriculum progressed into side control defence.'
    )
    doc.add_paragraph('The objective of this phase was to help students understand:')
    add_bullets(doc, [
        'How to survive underneath side control',
        'How to stay safe under pressure',
        'How to recover defensive structure',
        'How to prevent submissions and mount',
        'How to escape without panic',
    ])
    doc.add_paragraph(
        'The progression was designed so students could now understand the opposite side of the position '
        'they had previously learned from the top.'
    )

    add_heading(doc, 'Position Identity (Bottom)', 3)
    doc.add_paragraph('Students were introduced to the idea that defensive side control is based on:')
    add_bullets(doc, [
        'Frames',
        'Elbows and knees protecting space',
        'Controlling distance underneath pressure',
        'Reconnecting structure when pressure breaks it',
        'Patience and timing',
    ])
    doc.add_paragraph('The main objective from bottom side control is to:')
    add_bullets(doc, [
        'Stay safe',
        'Avoid submissions',
        'Avoid mount',
        'Recover structure',
        'Create opportunities to escape',
    ])
    doc.add_paragraph('Students were encouraged to understand that:')
    add_bullets(doc, [
        'Survival comes before escaping',
        'Recovering frames comes before movement',
        'Patience is more important than rushing escapes',
    ])

    add_heading(doc, 'Structure and Execution (Bottom)', 3)
    doc.add_paragraph('The progression started with:')
    add_bullets(doc, [
        'Understanding frames',
        'Protecting the elbows and hands',
        'Blocking mount with the knee',
        'Understanding how space is created and lost underneath pressure',
    ])
    doc.add_paragraph('Before moving into:')
    add_bullets(doc, [
        'Recovering frames',
        'Reconnecting defensive structure',
        'Escaping and recovering guard',
    ])
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'How frames protect space',
        'How body positioning affects survival',
        'How structure must be rebuilt before escaping',
    ])

    add_heading(doc, 'Week 1 — Survival and Frames', 3)
    doc.add_paragraph('Main focus:')
    add_bullets(doc, [
        'Keeping and recovering frames',
        'Hiding hands to avoid submissions',
        'Using the knee to block mount',
        'Understanding the relationship between hands, elbows and knees',
        'Recovering structure when frames are lost',
    ])
    doc.add_paragraph('Main goal: stay safe under pressure without getting mounted or submitted easily.')
    doc.add_paragraph('Specific training included:')
    add_bullets(doc, [
        'Top player working to clear frames, separate elbows, switch bases and progress to mount',
        'Bottom player working to maintain and recover defensive structure',
    ])
    doc.add_paragraph('This also reinforced concepts previously learned during the side control top block.')
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'How top pressure works',
        'How frames slow progression',
        'How defensive structure creates opportunities to survive and eventually escape',
    ])

    # --- REACT (Op) — Escapes ---
    add_heading(doc, 'React (Op) — Escapes', 2)
    doc.add_paragraph('From midweek onwards, progression moved into escapes.')
    doc.add_paragraph('Main focus:')
    add_bullets(doc, [
        'Recovering guard (main option)',
        'Turtle as secondary option',
    ])
    doc.add_paragraph('Important concepts:')
    add_bullets(doc, [
        'Every escape creates exposure for attacks',
        'If the escape fails, students should reset to strong defensive structure',
        'Escaping requires patience and timing',
        'Panic usually creates more openings for the top player',
    ])
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'How to recover defensive structure before trying again',
        'How to wait for movement and attacking opportunities from the top player',
        'How survival and escape work together',
    ])
    doc.add_paragraph('Main message throughout the block: defence comes first, attack comes later.')
    doc.add_paragraph('The objective of the block was to help students develop:')
    add_bullets(doc, [
        'Confidence under pressure',
        'Patience',
        'Defensive awareness',
        'Long-term retention through feeling safer during sparring',
    ])
    doc.add_paragraph(
        'The curriculum emphasised that students with strong defensive understanding will usually '
        'improve faster, feel less frustration, stay calmer during sparring, and remain more motivated long term.'
    )
    doc.add_paragraph(
        'Once students understood defensive structure and survival concepts, they began connecting '
        'bottom side control to recovering guard, rebuilding frames, turtle transitions, '
        'creating space, and timing escapes with opponent movement.'
    )
    doc.add_paragraph('Key concepts:')
    add_bullets(doc, [
        'Survival before escape',
        'Recover frames before movement',
        'Patience before explosiveness',
    ])

    # ══════════════════════════════════════════════════════════════
    # LEVEL 0: PASSED — MOUNT
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'Level 0 — Passed: Mount', 1)

    doc.add_paragraph(
        'After working on side control top and side control escapes, '
        'the curriculum progressed into mount (top position).'
    )
    doc.add_paragraph('The objective of this phase was to help students understand:')
    add_bullets(doc, [
        'What defines strong mount control',
        'How to maintain balance while mounted',
        'How to follow movement and reactions from the bottom player',
        'How mount progression creates stronger attacking opportunities',
        'How control leads to submissions',
    ])
    doc.add_paragraph(
        'The progression was designed so students could now begin understanding mount '
        'as a dominant control position rather than only an attacking position.'
    )

    # --- DEFINE ---
    add_heading(doc, 'Define', 2)
    doc.add_paragraph('Students were introduced to the idea that mount is based on:')
    add_bullets(doc, [
        'Balance',
        'Controlling movement and space',
        'Weight distribution',
        'Hip connection',
        'Following the partner\'s movement',
        'Maintaining control while the bottom player bridges and moves',
    ])
    doc.add_paragraph('The main objective of mount is to:')
    add_bullets(doc, [
        'Maintain dominant positioning',
        'Control the hips and shoulders',
        'Limit escape movement',
        'Progressively improve control',
        'Create attacking opportunities safely',
    ])
    doc.add_paragraph('Students were encouraged to understand that:')
    add_bullets(doc, [
        'Maintaining mount is more important than rushing submissions',
        'Balance and positioning are more important than squeezing or forcing control',
        'Mount progression creates stronger attacks and better control',
    ])

    # --- CONTROL (Me) ---
    add_heading(doc, 'Control (Me) — Mount Progression', 2)
    doc.add_paragraph('The progression started with:')
    add_bullets(doc, [
        'Low mount control',
        'Maintaining balance against bridging and movement',
        'Understanding how to follow reactions from the bottom player',
    ])
    doc.add_paragraph('Before moving into:')
    add_bullets(doc, [
        'Technical mount when the partner turns sideways',
        'Progressing into high mount',
        'Opening elbows and frames',
        'Improving upper body control',
    ])
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'How movement from the bottom creates transitions for the top player',
        'How mount progression improves control',
        'How high mount creates stronger attacking opportunities',
    ])

    add_heading(doc, 'Week 1 — Control and Mount Progression', 3)
    doc.add_paragraph('Main focus:')
    add_bullets(doc, [
        'Low mount control',
        'Balance and posture',
        'Maintaining the position against bridging and movement',
        'Transitioning to technical mount',
        'Progressing into high mount',
        'Opening elbows and frames',
        'Improving control before attacking',
    ])
    doc.add_paragraph(
        'Submissions were introduced during this phase, but the emphasis remained on '
        'maintaining mount, following movement, and understanding control before attack.'
    )
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'How reactions from the bottom player create transitions',
        'How mount progression strengthens attacks',
        'How losing balance usually means losing the position',
    ])

    # --- ATTACK (Me) ---
    add_heading(doc, 'Attack (Me) — Submissions From Mount', 2)

    add_heading(doc, 'Week 2 — Attacks From Mount', 3)
    doc.add_paragraph(
        'Once students became more comfortable maintaining and progressing the mount, '
        'focus shifted more towards submissions while maintaining control.'
    )
    doc.add_paragraph('Main focus:')
    add_bullets(doc, [
        'Maintaining mount while attacking',
        'Controlling posture during submissions',
        'Understanding reactions and defensive movement',
        'Connecting mount control with attacks',
    ])
    doc.add_paragraph('Main actions:')
    add_bullets(doc, [
        'Armbar from mount',
        'Americana from mount',
    ])
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'How submissions come from control and progression',
        'How high mount creates attacking opportunities',
        'How reactions from the bottom player open attacks and transitions',
    ])
    doc.add_paragraph(
        'Once students understood mount control and progression, they began connecting the position to '
        'armbar attacks, americana attacks, technical mount transitions, high mount progression, '
        'and maintaining balance while attacking.'
    )
    doc.add_paragraph('Key concepts:')
    add_bullets(doc, [
        'Control/balance before attack',
        'Progression before submission',
        'Maintain control while attacking',
    ])

    # --- DEFEND (Op) — Mount Defence ---
    add_heading(doc, 'Defend (Op) — Mount Defence', 2)
    doc.add_paragraph(
        'After working on mount control and mount attacks, the curriculum progressed into '
        'defending the mount position.'
    )
    doc.add_paragraph('The objective of this phase was to help students understand:')
    add_bullets(doc, [
        'How to survive underneath the mount',
        'How to stay calm under pressure',
        'How to use frames and hip movement effectively',
        'How to avoid being controlled flat on the back',
        'How to create transitions into safer positions',
    ])
    doc.add_paragraph(
        'The progression was designed so students could now understand the opposite side of the mount '
        'position they had previously learned from the top.'
    )

    add_heading(doc, 'Position Identity (Bottom)', 3)
    doc.add_paragraph('Students were introduced to the idea that defensive mount is based on:')
    add_bullets(doc, [
        'Staying calm and patient',
        'Avoiding flat positioning',
        'Hip engagement and movement',
        'Frames and structure',
        'Creating angles and transitions',
    ])
    doc.add_paragraph('The main objective from bottom mount is to:')
    add_bullets(doc, [
        'Survive',
        'Avoid strong control and submissions',
        'Create movement and instability',
        'Recover guard or safer positions',
    ])
    doc.add_paragraph('Students were encouraged to understand that:')
    add_bullets(doc, [
        'Staying calm is essential for escaping',
        'Movement and structure work together',
        'Lying flat gives the top player more control',
        'Good defensive structure creates escape opportunities',
    ])

    add_heading(doc, 'Structure and Execution (Bottom)', 3)
    doc.add_paragraph('The progression started with:')
    add_bullets(doc, [
        'Staying relaxed under pressure',
        'Flattening one leg and bridging',
        'Getting onto the side',
        'Avoiding underhooks and strong upper body control',
        'Understanding how frames create space',
    ])
    doc.add_paragraph('Before moving into:')
    add_bullets(doc, [
        'Guard recovery',
        'Half guard recovery',
        'Upa escape',
        'Recovering guard during side transitions',
    ])
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'How bridging disrupts balance',
        'How hip movement creates space',
        'How frames and movement must work together',
        'How staying active prevents the top player from stabilising',
    ])

    add_heading(doc, 'Key Principles', 3)
    doc.add_paragraph('Main concepts taught throughout the block:')
    add_bullets(doc, [
        'Stay calm and relaxed',
        'Stay on the side instead of flat on the back',
        'Engage the hips actively',
        'Establish and recover frames',
        'Stay active without exposing submissions',
    ])
    doc.add_paragraph('Students were consistently reminded not to panic, not to force escapes recklessly, '
                      'and to prioritise structure and timing.')

    # --- REACT (Op) — Mount Escapes ---
    add_heading(doc, 'React (Op) — Mount Escapes', 2)

    add_heading(doc, 'Defence 1 — Upa (Bridge and Roll)', 3)
    doc.add_paragraph('Students learned:')
    add_bullets(doc, [
        'Controlling one of the opponent\'s arms',
        'Trapping the arm at the elbow and shoulder',
        'Stepping over the opponent\'s leg on the same side',
        'Bridging and rolling using hip drive and elbow pressure',
    ])
    doc.add_paragraph('Main concepts:')
    add_bullets(doc, [
        'Breaking balance',
        'Trapping posts',
        'Using hip power rather than upper body strength',
    ])
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'How bridging creates instability',
        'How removing posts creates openings for reversal',
    ])

    add_heading(doc, 'Defence 2 — Recover Guard / Half Guard', 3)
    doc.add_paragraph('Main focus:')
    add_bullets(doc, [
        'Staying on the side',
        'Framing at the hips and knee',
        'Creating space using elbows and hips',
        'Recovering guard or half guard safely',
    ])
    doc.add_paragraph('Students learned:')
    add_bullets(doc, [
        'One elbow framing at the hips',
        'The opposite arm controlling inside the knee',
        'Combining frames with hip movement',
        'Sliding the legs back inside to recover position',
    ])
    doc.add_paragraph('Main concepts:')
    add_bullets(doc, [
        'Structure before movement',
        'Frames create space',
        'Hips reconnect the guard',
    ])

    add_heading(doc, 'Defence 3 — Recover Guard During Side Transition', 3)
    doc.add_paragraph(
        'Students also worked on situations where the opponent transitioned towards '
        'side control from mount.'
    )
    doc.add_paragraph('Main focus:')
    add_bullets(doc, [
        'Maintaining side positioning',
        'Keeping active frames',
        'Controlling the knee and hips',
        'Preventing the opponent from settling into side control',
    ])
    doc.add_paragraph('Students learned:')
    add_bullets(doc, [
        'Framing against the hips',
        'Controlling the knee with the same arm',
        'Gripping the bottom of the trousers with the bottom hand',
        'Using hip movement to create space and replace guard',
    ])
    doc.add_paragraph('Main concepts:')
    add_bullets(doc, [
        'Transitions create escape opportunities',
        'Staying active prevents stabilisation',
        'Reconnecting structure before fully escaping',
    ])

    doc.add_paragraph(
        'Once students understood defensive structure from mount, they began connecting the position to '
        'guard recovery, half guard recovery, reversals, escaping during transitions, '
        'and rebuilding defensive structure.'
    )
    doc.add_paragraph('Key concepts:')
    add_bullets(doc, [
        'Stay calm before escaping',
        'Frames before movement',
        'Hips create space',
        'Structure before explosiveness',
    ])

    # ══════════════════════════════════════════════════════════════
    # LEVEL 0: PASSED — BACK CONTROL
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'Level 0 — Passed: Back Control', 1)

    doc.add_paragraph(
        'After working on mount control, mount attacks and mount escapes, '
        'the curriculum progressed into back control and attacks.'
    )
    doc.add_paragraph('The objective of this phase was to help students understand:')
    add_bullets(doc, [
        'What defines strong back control',
        'How to maintain chest-to-back connection',
        'How to use the seat belt and hooks correctly',
        'How to follow movement and prevent escapes',
        'How control creates attacking and transitional opportunities',
    ])
    doc.add_paragraph(
        'Students were introduced to the idea that back control is one of the strongest attacking '
        'positions in Jiu-Jitsu, but also one of the most difficult dominant positions to maintain.'
    )

    # --- DEFINE ---
    add_heading(doc, 'Define', 2)
    doc.add_paragraph('Students were introduced to the idea that back control is based on:')
    add_bullets(doc, [
        'Chest-to-back connection',
        'Seat belt control',
        'Hooks and leg engagement',
        'Controlling upper body movement',
        'Following the opponent\'s movement and reactions',
    ])
    doc.add_paragraph('The main objective of back control is to:')
    add_bullets(doc, [
        'Maintain connection to the opponent\'s back',
        'Prevent the opponent escaping or turning',
        'Control movement using upper and lower body together',
        'Create attacking opportunities safely',
    ])
    doc.add_paragraph('Students were encouraged to understand that:')
    add_bullets(doc, [
        'The seat belt is the main control of the position',
        'The hooks assist control and movement',
        'Chest connection is more important than squeezing',
        'The head position helps maintain alignment and control',
    ])

    # --- CONTROL (Me) ---
    add_heading(doc, 'Control (Me)', 2)
    doc.add_paragraph('The progression started with:')
    add_bullets(doc, [
        'Understanding seat belt control',
        'Understanding the role of the hooks',
        'Maintaining chest-to-back connection',
        'Keeping the head positioned over the underhook side',
        'Following movement with the hips and legs',
    ])
    doc.add_paragraph('Students first worked on:')
    add_bullets(doc, [
        'Simply holding and maintaining the back position',
        'Preventing escapes',
        'Understanding how movement from the bottom affects control',
    ])
    doc.add_paragraph('Specific training included:')
    add_bullets(doc, [
        'Students attempting to maintain back control while the partner tried to escape',
        'Stopping and regrouping after rounds to discuss questions and observations',
    ])
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'How difficult back control is to maintain',
        'How the legs must stay active during the position',
        'How the seat belt keeps the opponent connected to the chest',
    ])

    add_heading(doc, 'Main Concepts', 3)
    doc.add_paragraph('Key concepts reinforced throughout the block:')
    add_bullets(doc, [
        'Maintain chest-to-back connection',
        'Use the seat belt to control upper body movement',
        'Keep the head over the underhook side',
        'Work actively with the hooks and legs',
        'Follow movement instead of squeezing too hard',
    ])
    doc.add_paragraph('Students were reminded that:')
    add_bullets(doc, [
        'Losing chest connection usually means losing the back',
        'The hooks help follow movement and maintain positioning',
        'Control comes before submissions',
    ])

    # --- ATTACK (Me) ---
    add_heading(doc, 'Attack (Me) — Submissions', 2)
    doc.add_paragraph(
        'Once students became more comfortable maintaining the position, submissions were introduced.'
    )
    doc.add_paragraph('Main submissions:')
    add_bullets(doc, [
        'Collar choke',
        'Rear naked choke',
    ])
    doc.add_paragraph('The focus remained on:')
    add_bullets(doc, [
        'Maintaining connection during attacks',
        'Controlling reactions before attacking',
        'Using the seat belt and chest connection to create openings safely',
    ])
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'Submissions come from strong control',
        'Chest connection and head positioning assist the attacks',
        'The opponent\'s defensive reactions create attacking opportunities',
    ])

    # --- TRANSIT (Me) ---
    add_heading(doc, 'Transit (Me) — Transitions', 2)
    doc.add_paragraph('Students were also introduced to the idea that:')
    add_bullets(doc, [
        'If the opponent places the back on the floor and escapes the angle of control, '
        'the top player should transition to mount rather than forcing the back position.',
    ])
    doc.add_paragraph('This helped students understand:')
    add_bullets(doc, [
        'How dominant positions connect together',
        'How transitions maintain control during escapes and reactions',
    ])

    doc.add_paragraph(
        'Once students understood the structure and control of back control, '
        'they began connecting the position to maintaining chest connection, '
        'following movement with hooks and legs, preventing escapes, '
        'collar choke attacks, rear naked choke attacks, and transitions to mount.'
    )
    doc.add_paragraph('Key concepts:')
    add_bullets(doc, [
        'Chest connection prevents escapes',
        'Hand fight to achieve submissions',
        'Transition before losing control',
    ])

    # ══════════════════════════════════════════════════════════════
    # LEVEL 0: PASSED — TURTLE
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'Level 0 — Passed: Turtle', 1)

    doc.add_paragraph(
        'After working on back control and mount defence, the curriculum progressed into the turtle position.'
    )
    doc.add_paragraph('The objective of this phase was to help students understand:')
    add_bullets(doc, [
        'What turtle position is',
        'When turtle should be used',
        'How to stay safe in turtle',
        'How to avoid back takes and attacks',
        'How to transition back to safer positions',
    ])
    doc.add_paragraph(
        'Students were introduced to the idea that turtle is mainly a defensive transition position '
        'used to avoid side control, avoid the guard pass, and recover position and movement.'
    )
    doc.add_paragraph(
        'Students were also taught that turtle is not considered a guard position in this curriculum because '
        'there are no direct sweeps from the position — '
        'the main purpose is defensive recovery and transition rather than attacking.'
    )

    # --- DEFINE ---
    add_heading(doc, 'Define', 2)
    doc.add_paragraph('Students were introduced to the idea that turtle position is based on:')
    add_bullets(doc, [
        'Strong defensive structure',
        'Protecting the neck and back',
        'Preventing hooks and back control',
        'Staying compact and connected',
    ])
    doc.add_paragraph('The main objective of turtle is to:')
    add_bullets(doc, [
        'Avoid dominant control',
        'Prevent attacks and back takes',
        'Create opportunities to recover guard or reverse position',
    ])
    doc.add_paragraph('Students were encouraged to understand that:')
    add_bullets(doc, [
        'Turtle should remain active, not passive',
        'Structure and positioning are more important than strength',
        'Exposing the neck or elbows creates openings for attacks',
    ])

    # --- CONTROL ---
    add_heading(doc, 'Control — Defensive Structure', 2)
    doc.add_paragraph('The progression started with:')
    add_bullets(doc, [
        'Both knees wide on the floor',
        'Head positioned on the floor',
        'Elbows connected near the lap/hips',
        'Hands protecting and closing near the neck',
    ])
    doc.add_paragraph('Main concepts:')
    add_bullets(doc, [
        'Keeping the elbows tight to avoid back takes',
        'Protecting the neck and collar area',
        'Staying compact and balanced',
        'Using structure instead of muscling the position',
    ])
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'How elbow positioning prevents hooks and control',
        'How posture affects defensive safety',
        'How movement and timing create recovery opportunities',
    ])

    add_heading(doc, 'Specific Training', 3)
    doc.add_paragraph(
        'Specific sparring was used throughout the block to help students apply '
        'the defensive concepts under pressure.'
    )
    doc.add_paragraph('Main objectives:')
    add_bullets(doc, [
        'Top player works to establish control or expose the back',
        'Bottom player works to maintain structure and recover position',
    ])
    doc.add_paragraph('Students started from turtle position with:')
    add_bullets(doc, [
        'Top player attempting to establish back control, hooks or transitions',
        'Bottom player attempting to stay compact, protect the neck and recover guard or reverse position',
    ])
    doc.add_paragraph('The focus remained on:')
    add_bullets(doc, [
        'Maintaining structure under pressure',
        'Staying calm while moving',
        'Protecting the elbows and neck',
        'Using timing instead of explosiveness',
    ])
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'Turtle is a transitional position',
        'Movement and reactions create recovery opportunities',
        'Remaining passive usually allows the top player to stabilise control',
    ])

    # --- TRANSIT (Me) — Actions ---
    add_heading(doc, 'Transit (Me) — Actions From Turtle', 2)
    doc.add_paragraph(
        'Once students understood the structure and defensive concepts of turtle, '
        'they began connecting the position to recovery and reversal actions.'
    )
    doc.add_paragraph('Main actions:')
    add_bullets(doc, [
        'Replacing guard',
        'Rolling over the opponent',
    ])

    add_heading(doc, 'Technique 1 — Replace Guard', 3)
    doc.add_paragraph('Main focus:')
    add_bullets(doc, [
        'Maintaining defensive structure',
        'Creating space with movement and frames',
        'Reconnecting the legs and hips',
        'Recovering guard safely',
    ])
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'Turtle is a transition position',
        'The objective is to recover safer positions rather than remain static',
    ])

    add_heading(doc, 'Technique 2 — Roll Over the Opponent', 3)
    doc.add_paragraph('Main focus:')
    add_bullets(doc, [
        'Using timing and movement',
        'Redirecting pressure',
        'Creating reversal opportunities from turtle transitions',
    ])
    doc.add_paragraph('Students were encouraged to understand:')
    add_bullets(doc, [
        'Movement from the top player creates opportunities',
        'Timing and structure are more important than explosiveness',
    ])

    add_heading(doc, 'Main Concepts', 3)
    doc.add_paragraph('Key concepts reinforced throughout the block:')
    add_bullets(doc, [
        'Turtle is defensive, not offensive',
        'Elbows protect against back takes',
        'Neck protection is essential',
        'Movement should stay compact and controlled',
        'Turtle should connect to guard recovery or reversals',
    ])
    doc.add_paragraph(
        'The objective of the block was to help students understand turtle as '
        'a defensive transition position first, then a platform for recovering '
        'safer positions or reversing movement.'
    )
    doc.add_paragraph('Key concepts:')
    add_bullets(doc, [
        'Structure before movement',
        'Protection before escape',
        'Transition before staying static',
    ])

    return doc


if __name__ == '__main__':
    doc = build_doc()
    out = '/home/ubuntu/BBJJ/Curriculum_Fundamentals_Rewrite.docx'
    doc.save(out)
    print(f'Saved → {out}')
